# File: backend/app/services/knowledge_service.py
#
# RAG pipeline — ChromaDB + sentence-transformers
# Install: pip install chromadb sentence-transformers pypdf2

import os
# Must be set before sentence-transformers / transformers import.
# TensorFlow has a broken protobuf dependency in this environment;
# forcing the PyTorch backend avoids the crash.
os.environ.setdefault("USE_TF", "0")
os.environ.setdefault("USE_TORCH", "1")

import logging
import re
import uuid
import json
import threading
import time
import concurrent.futures
from datetime import datetime

logger = logging.getLogger(__name__)
from typing import Optional, List

CHROMA_DIR     = os.path.join(os.path.dirname(__file__), "..", "..", "data", "knowledge_base")
DOCS_META_PATH = os.path.join(CHROMA_DIR, "documents_meta.json")
os.makedirs(CHROMA_DIR, exist_ok=True)

# ── Singletons ────────────────────────────────────────────────────────────────
_collection       = None
_embedder         = None
_anthropic_client = None


def _get_anthropic_client():
    global _anthropic_client
    if _anthropic_client is None:
        import anthropic
        from app.core.config import settings
        _anthropic_client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
    return _anthropic_client


def _get_collection():
    global _collection
    if _collection is None:
        try:
            import chromadb
            client = chromadb.PersistentClient(path=CHROMA_DIR)
            _collection = client.get_or_create_collection(
                name="nexusdesk_knowledge",
                metadata={"hnsw:space": "cosine"},
            )
        except ImportError:
            raise RuntimeError("chromadb not installed. Run: pip install chromadb")
    return _collection


def _get_embedder():
    global _embedder
    if _embedder is None:
        try:
            import os
            # TensorFlow has a broken protobuf dependency in this environment.
            # Force transformers to use the PyTorch backend only.
            os.environ.setdefault("USE_TF", "0")
            os.environ.setdefault("USE_TORCH", "1")
            from sentence_transformers import SentenceTransformer
            _embedder = SentenceTransformer("all-MiniLM-L6-v2")
        except ImportError:
            raise RuntimeError("sentence-transformers not installed.")
    return _embedder


# ── Text extraction ───────────────────────────────────────────────────────────

def _clean_pdf_text(text: str) -> str:
    """
    Clean raw PDF-extracted text to restore readable structure.
    PDF extraction often produces flat text with no paragraph breaks.
    """
    # Normalize whitespace
    text = text.replace("\t", "  ")
    text = re.sub(r" {3,}", "  ", text)

    # Remove standalone page numbers
    text = re.sub(r"(?m)^\d+$", "", text)

    # Normalize bullet characters
    text = text.replace("●", "•").replace("◆", "•").replace("▪", "•")

    # Detect flat PDF text: average line length > 150 chars means structure was lost
    lines     = [l for l in text.split("\n") if l.strip()]
    avg_len   = sum(len(l) for l in lines) / max(len(lines), 1)

    if avg_len > 150:
        # Split before numbered items: "something. 1. Next item"
        text = re.sub(r"([.!?])\s+(\d+[.:]\s+[A-Z])", r"\1\n\2", text)
        # Split before bullets after sentence endings
        text = re.sub(r"([.!?])\s+(•\s)", r"\1\n\2", text)
        # Split on common section keywords
        section_kw = (
            r"Introduction|Summary|Overview|Conclusion|Background|Components|"
            r"Architecture|Troubleshooting|Resolution|Root Cause|Preventive|"
            r"Key Takeaway|End of Document|Section|Appendix|References|"
            r"Purpose|Scope|High-Level|Useful|Important"
        )
        text = re.sub(
            rf"([.!?])\s+((?:{section_kw})\s*[\d.:)]*\s)",
            r"\1\n\n\2",
            text,
        )
        # Split ALL CAPS headings out of running text
        text = re.sub(
            r"([a-z.!?])\s+([A-Z]{3,}(?:\s+[A-Z]{2,}){0,4})\s*:",
            r"\1\n\n\2:",
            text
        )

    # Clean up excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_docx_text(content: bytes) -> str:
    """Extract text from .docx preserving headings and structure."""
    try:
        import io
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            style = para.style.name if para.style else ""
            if "Heading 1" in style:
                lines.append(f"\n{text.upper()}")
            elif "Heading 2" in style:
                lines.append(f"\n{text}")
            else:
                lines.append(text)
        # Extract tables
        for table in doc.tables:
            lines.append("")
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    lines.append("  |  ".join(row_cells))
            lines.append("")
        text = "\n".join(lines)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
    except Exception as e:
        logger.warning("[KB] DOCX extraction error: %s", e)
        return ""


def _extract_text(content: bytes, filename: str) -> str:
    """Extract text from file, preserving structure."""
    fname = filename.lower()

    if fname.endswith(".docx"):
        return _extract_docx_text(content)

    if fname.endswith(".pdf"):
        for lib in ["PyPDF2", "pypdf"]:
            try:
                mod = __import__(lib)
                import io
                reader = mod.PdfReader(io.BytesIO(content))
                pages = []
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages.append(page_text.strip())
                raw = "\n\n".join(pages)
                return _clean_pdf_text(raw)
            except (ImportError, Exception):
                continue

    # Plain text / markdown — decode as-is
    return content.decode("utf-8", errors="ignore")


def _chunk_text(text: str, chunk_size: int = 200, overlap: int = 30) -> List[str]:
    """
    Chunk by paragraphs first, then by size — preserving line breaks and structure.
    Never joins paragraphs into a flat blob.
    """
    # Split into paragraphs
    paragraphs = re.split(r"\n{2,}", text)
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    chunks        = []
    current_paras = []
    current_len   = 0

    for para in paragraphs:
        para_words = para.split()
        para_len   = len(para_words)

        # Paragraph too large on its own — split by sentences
        if para_len > chunk_size:
            # First flush current buffer
            if current_paras:
                chunks.append("\n\n".join(current_paras))
                current_paras = []
                current_len   = 0
            # Split large paragraph into sentence chunks
            sentences = re.split(r"(?<=[.!?])\s+", para)
            sent_buf  = []
            sent_len  = 0
            for sent in sentences:
                sw = len(sent.split())
                if sent_len + sw > chunk_size and sent_buf:
                    chunks.append(" ".join(sent_buf))
                    # Keep last sentence as overlap
                    sent_buf  = [sent_buf[-1], sent] if sent_buf else [sent]
                    sent_len  = len(" ".join(sent_buf).split())
                else:
                    sent_buf.append(sent)
                    sent_len += sw
            if sent_buf:
                chunks.append(" ".join(sent_buf))
            continue

        # Adding this paragraph would exceed chunk_size
        if current_len + para_len > chunk_size and current_paras:
            chunks.append("\n\n".join(current_paras))
            # Overlap: keep last paragraph for context continuity
            overlap_para  = current_paras[-1] if current_paras else ""
            current_paras = [overlap_para, para] if overlap_para else [para]
            current_len   = len(overlap_para.split()) + para_len
        else:
            current_paras.append(para)
            current_len += para_len

    if current_paras:
        chunks.append("\n\n".join(current_paras))

    return [c.strip() for c in chunks if c.strip()]


# ── AI Summary Generation ─────────────────────────────────────────────────────

def _generate_summary(text: str, title: str, domain: str) -> str:
    """Generate a concise AI summary of the document for engineers."""
    try:
        client = _get_anthropic_client()

        prompt = f"""You are summarizing an IT support knowledge base document for engineers.

Document title: {title}
Domain: {domain}

Document content:
{text}

Write a thorough, structured summary for engineers. Include:
1. What this document covers (2-3 sentences)
2. Key technical concepts explained (4-6 bullet points with enough detail to be useful)
3. All diagnostic steps or commands mentioned (as bullet points)
4. Resolution approaches covered (as bullet points)
5. When to use this document (1-2 sentences)

Aim for 300-400 words. Be specific — include actual commands, IP ranges, thresholds, and technical details from the document. Use plain text only, no markdown headers or bold."""

        resp = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}],
        )
        return resp.content[0].text.strip()
    except Exception as e:
        logger.warning("[KB] Summary generation failed: %s", e)
        # Fallback: use first 300 chars of cleaned text
        return text[:300].strip() + "..."


# ── Metadata store ─────────────────────────────────────────────────────────────
# Protects documents_meta.json from concurrent writes by request threads and
# the background _auto_refresh_loop thread.
_meta_lock = threading.Lock()

# Serialises ChromaDB writes (add/delete) so concurrent uploads don't corrupt the index
_collection_write_lock = threading.Lock()

# RAG query result cache — keyed by query hash, value is (context_str, monotonic_timestamp)
_rag_cache: dict        = {}
_rag_cache_lock         = threading.Lock()
_RAG_CACHE_TTL          = 300   # seconds — invalidate after 5 minutes
_RAG_CACHE_MAX          = 500   # max entries before LRU eviction

# Thread pool for running blocking ChromaDB calls with a hard timeout
_chroma_executor        = concurrent.futures.ThreadPoolExecutor(max_workers=4, thread_name_prefix="chroma")
_CHROMA_TIMEOUT         = 2.0   # seconds


def _chroma_with_timeout(fn, *args, **kwargs):
    """Run a ChromaDB call in the thread pool with a hard 2s timeout. Returns None on timeout."""
    future = _chroma_executor.submit(fn, *args, **kwargs)
    try:
        return future.result(timeout=_CHROMA_TIMEOUT)
    except concurrent.futures.TimeoutError:
        logger.warning("[ChromaDB] Query timed out after %.1fs — returning empty context", _CHROMA_TIMEOUT)
        return None
    except Exception:
        raise


def _rag_cache_key(query: str, domain, n_results: int) -> str:
    import hashlib
    return hashlib.md5(f"{query}|{domain}|{n_results}".encode(), usedforsecurity=False).hexdigest()


def _clear_rag_cache() -> None:
    """Invalidate the entire RAG cache (called after any document add/delete)."""
    with _rag_cache_lock:
        _rag_cache.clear()


def _load_meta() -> dict:
    with _meta_lock:
        if os.path.exists(DOCS_META_PATH):
            with open(DOCS_META_PATH) as f:
                return json.load(f)
        return {}


def _save_meta(meta: dict):
    with _meta_lock:
        with open(DOCS_META_PATH, "w") as f:
            json.dump(meta, f, indent=2)


# ── Upload ────────────────────────────────────────────────────────────────────

def _extract_notion_page_id(url: str) -> str:
    """Extract Notion page ID from any Notion URL format."""
    import re as _re
    # Match 32-char hex at end of path (before ? or end)
    match = _re.search(r"([0-9a-f]{32})(?:\?|$)", url)
    if match:
        raw = match.group(1)
        return f"{raw[:8]}-{raw[8:12]}-{raw[12:16]}-{raw[16:20]}-{raw[20:]}"
    # Already formatted with dashes
    match = _re.search(r"([0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12})", url)
    if match:
        return match.group(1)
    raise ValueError(f"Could not extract Notion page ID from URL: {url}")


def _fetch_notion_content(url: str) -> str:
    """Fetch content from Notion page using the Notion API."""
    import urllib.request
    import json as _json

    from app.core.config import settings
    token = getattr(settings, "NOTION_API_TOKEN", "") or ""
    if not token:
        raise ValueError("NOTION_API_TOKEN not configured. Add it to .env to fetch Notion pages.")

    page_id = _extract_notion_page_id(url)

    headers_api = {
        "Authorization": f"Bearer {token}",
        "Notion-Version": "2022-06-28",
        "Content-Type": "application/json",
    }

    def api_get(endpoint: str) -> dict:
        req = urllib.request.Request(
            f"https://api.notion.com/v1/{endpoint}",
            headers=headers_api,
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            return _json.loads(resp.read().decode("utf-8"))

    def api_post(endpoint: str, body: dict) -> dict:
        data = _json.dumps(body).encode("utf-8")
        req = urllib.request.Request(
            f"https://api.notion.com/v1/{endpoint}",
            data=data,
            headers=headers_api,
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            return _json.loads(resp.read().decode("utf-8"))

    def extract_rich_text(rich_text_list: list) -> str:
        return "".join(rt.get("plain_text", "") for rt in rich_text_list)

    def block_to_text(block: dict) -> str:
        btype = block.get("type", "")
        data  = block.get(btype, {})
        lines = []

        if btype in ("paragraph", "quote", "callout"):
            text = extract_rich_text(data.get("rich_text", []))
            if text.strip():
                lines.append(text)

        elif btype in ("heading_1", "heading_2", "heading_3"):
            text = extract_rich_text(data.get("rich_text", []))
            prefix = {"heading_1": "# ", "heading_2": "## ", "heading_3": "### "}.get(btype, "")
            if text.strip():
                lines.append(f"\n{prefix}{text}")

        elif btype in ("bulleted_list_item", "numbered_list_item", "to_do"):
            text = extract_rich_text(data.get("rich_text", []))
            if text.strip():
                lines.append(f"• {text}")

        elif btype == "code":
            text = extract_rich_text(data.get("rich_text", []))
            lang = data.get("language", "")
            if text.strip():
                lines.append(f"```{lang}\n{text}\n```")

        elif btype == "table_row":
            cells = data.get("cells", [])
            row   = " | ".join(extract_rich_text(cell) for cell in cells)
            if row.strip():
                lines.append(row)

        elif btype == "divider":
            lines.append("---")

        return "\n".join(lines)

    # Fetch all blocks with pagination
    all_blocks = []
    cursor     = None
    while True:
        params = f"page_size=100"
        if cursor:
            params += f"&start_cursor={cursor}"
        result = api_get(f"blocks/{page_id}/children?{params}")
        all_blocks.extend(result.get("results", []))
        if not result.get("has_more"):
            break
        cursor = result.get("next_cursor")

    # Convert blocks to text
    text_parts = []
    for block in all_blocks:
        text = block_to_text(block)
        if text.strip():
            text_parts.append(text)

        # Fetch children for blocks that have them (tables, toggles etc.)
        if block.get("has_children"):
            try:
                children = api_get(f"blocks/{block['id']}/children?page_size=100")
                for child in children.get("results", []):
                    child_text = block_to_text(child)
                    if child_text.strip():
                        text_parts.append(child_text)
            except Exception:
                pass

    full_text = "\n\n".join(text_parts)
    full_text = re.sub(r"\n{3,}", "\n\n", full_text)
    return full_text.strip()


_SSRF_BLOCKED_NETS = [
    # IPv4 private/link-local/loopback
    "10.", "172.16.", "172.17.", "172.18.", "172.19.", "172.20.", "172.21.",
    "172.22.", "172.23.", "172.24.", "172.25.", "172.26.", "172.27.", "172.28.",
    "172.29.", "172.30.", "172.31.", "192.168.", "127.", "169.254.", "0.",
    # IPv6 private
    "::1", "fc", "fd", "fe80",
]


def _validate_url_for_fetch(url: str) -> None:
    """Raise ValueError if the URL could reach internal/SSRF-sensitive targets."""
    import urllib.parse
    import socket

    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError(f"URL scheme '{parsed.scheme}' is not allowed. Only http and https are permitted.")

    host = parsed.hostname or ""
    if not host:
        raise ValueError("URL has no hostname.")

    # Reject raw IPs that look internal — do a quick check before DNS
    for prefix in _SSRF_BLOCKED_NETS:
        if host.startswith(prefix):
            raise ValueError("URL resolves to a private/internal network address.")

    # Resolve hostname and check resolved IPs
    try:
        results = socket.getaddrinfo(host, None)
        for _, _, _, _, sockaddr in results:
            ip = sockaddr[0]
            for prefix in _SSRF_BLOCKED_NETS:
                if ip.startswith(prefix):
                    raise ValueError("URL resolves to a private/internal network address.")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Could not resolve hostname '{host}': {e}")


def fetch_url_content(url: str) -> str:
    """
    Fetch and extract clean text from any URL.
    Notion pages use the Notion API.
    All other URLs use web scraping.
    """
    import urllib.request
    import urllib.parse
    import html as html_module

    # Always validate scheme and block SSRF before any routing decision.
    # A URL like http://169.254.169.254/?notion.so must be rejected here,
    # not accidentally routed through the Notion path.
    _parsed = urllib.parse.urlparse(url)
    if _parsed.scheme not in ("http", "https"):
        raise ValueError(f"URL scheme '{_parsed.scheme}' is not allowed.")

    # Notion — use API (only if hostname is actually notion.so/notion.site)
    _host = (_parsed.hostname or "").lower()
    if _host == "notion.so" or _host.endswith(".notion.so") or \
       _host == "notion.site" or _host.endswith(".notion.site"):
        return _fetch_notion_content(url)

    # Block SSRF before making any network request
    _validate_url_for_fetch(url)

    # Generic web scraping for all other URLs
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
    }

    try:
        req  = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=15) as resp:
            raw_html = resp.read().decode("utf-8", errors="ignore")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Could not fetch URL: {e}")

    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw_html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "meta", "noscript"]):
            tag.decompose()
        main = soup.find("main") or soup.find("article") or soup.find("body")
        text = main.get_text(separator="\n", strip=True) if main else soup.get_text(separator="\n", strip=True)
    except ImportError:
        text = re.sub(r"<script[^>]*>.*?</script>", "", raw_html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = html_module.unescape(text)

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    text  = "\n".join(lines)
    text  = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def upload_url(
    url: str,
    title: str,
    domain: str,
    description: str,
    uploaded_by: str,
    uploaded_by_role: str = "admin",
) -> dict:
    """Fetch a URL and index its content into the knowledge base."""
    try:
        text = fetch_url_content(url)
    except ValueError as e:
        return {"success": False, "error": str(e)}

    if not text.strip():
        return {"success": False, "error": "Could not extract text from URL — page may be empty or require login"}

    chunks     = _chunk_text(text)
    embedder   = _get_embedder()
    embeddings = embedder.encode(chunks).tolist()
    collection = _get_collection()

    doc_id    = str(uuid.uuid4())
    ids       = [f"{doc_id}_{i}" for i in range(len(chunks))]
    metadatas = [
        {"doc_id": doc_id, "domain": domain, "title": title, "chunk_index": i, "source_url": url}
        for i in range(len(chunks))
    ]

    with _collection_write_lock:
        collection.add(ids=ids, embeddings=embeddings, documents=chunks, metadatas=metadatas)

    summary      = _generate_summary(text[:8000], title, domain)
    content_hash = _content_hash(text)

    meta = _load_meta()
    meta[doc_id] = {
        "id":               doc_id,
        "title":            title,
        "filename":         url,
        "source_url":       url,
        "domain":           domain,
        "description":      description,
        "uploaded_by":      uploaded_by,
        "uploaded_by_role": uploaded_by_role,
        "chunk_count":      len(chunks),
        "summary":          summary,
        "content_hash":     content_hash,
        "last_refreshed":   datetime.utcnow().isoformat(),
        "created_at":       datetime.utcnow().isoformat(),
        "source_type":      "url",
    }
    _save_meta(meta)
    _clear_rag_cache()

    logger.info("[KB] URL '%s' indexed — %d chunks [%s]", title, len(chunks), uploaded_by_role)
    return {
        "success":     True,
        "doc_id":      doc_id,
        "title":       title,
        "chunk_count": len(chunks),
        "summary":     summary,
        "message":     f"'{title}' indexed from URL — {len(chunks)} chunks.",
    }


def upload_document(
    content: bytes,
    filename: str,
    title: str,
    domain: str,
    description: str,
    uploaded_by: str,
    uploaded_by_role: str = "admin",
) -> dict:
    text = _extract_text(content, filename)
    if not text.strip():
        return {"success": False, "error": "Could not extract text from document"}

    chunks     = _chunk_text(text)
    embedder   = _get_embedder()
    embeddings = embedder.encode(chunks).tolist()
    collection = _get_collection()

    doc_id    = str(uuid.uuid4())
    ids       = [f"{doc_id}_{i}" for i in range(len(chunks))]
    metadatas = [
        {"doc_id": doc_id, "domain": domain, "title": title, "chunk_index": i}
        for i in range(len(chunks))
    ]

    with _collection_write_lock:
        collection.add(ids=ids, embeddings=embeddings, documents=chunks, metadatas=metadatas)

    # Generate AI summary
    summary = _generate_summary(text[:8000], title, domain)

    meta = _load_meta()
    meta[doc_id] = {
        "id":               doc_id,
        "title":            title,
        "filename":         filename,
        "domain":           domain,
        "description":      description,
        "uploaded_by":      uploaded_by,
        "uploaded_by_role": uploaded_by_role,
        "chunk_count":      len(chunks),
        "summary":          summary,
        "created_at":       datetime.utcnow().isoformat(),
    }
    _save_meta(meta)
    _clear_rag_cache()

    logger.info("[KB] Upload '%s' — %d chunks [%s]", title, len(chunks), uploaded_by_role)
    return {
        "success":     True,
        "doc_id":      doc_id,
        "title":       title,
        "chunk_count": len(chunks),
        "summary":     summary,
        "message":     f"'{title}' indexed — {len(chunks)} chunks.",
    }


# ── Search ────────────────────────────────────────────────────────────────────

def search_knowledge(
    query: str,
    n_results: int = 5,
    domain: Optional[str] = None,
) -> dict:
    try:
        collection = _get_collection()
        total = collection.count()
        if total == 0:
            return {"query": query, "results": [], "total": 0}

        embedder        = _get_embedder()
        query_embedding = embedder.encode([query]).tolist()
        where = {"domain": domain} if domain and domain not in ("", "other", "all") else None

        results = collection.query(
            query_embeddings=query_embedding,
            n_results=min(n_results, total),
            where=where,
            include=["documents", "metadatas", "distances"],
        )

        # Deduplicate by doc_id — keep only the best matching chunk per document
        # This prevents the same document showing up 7 times with different chunks
        seen_docs = {}
        meta = _load_meta()
        if results["documents"] and results["documents"][0]:
            for doc, metadata, dist in zip(
                results["documents"][0],
                results["metadatas"][0],
                results["distances"][0],
            ):
                doc_id   = metadata.get("doc_id", "")
                doc_meta = meta.get(doc_id, {})
                sim      = round((1 - dist) * 100, 1)

                # Only keep the highest similarity chunk per document
                if doc_id not in seen_docs or sim > seen_docs[doc_id]["cosine_similarity"]:
                    seen_docs[doc_id] = {
                        "content":           doc,
                        "title":             metadata.get("title", doc_meta.get("title", "Unknown")),
                        "doc_id":            doc_id,
                        "domain":            metadata.get("domain", "other"),
                        "cosine_similarity": sim,
                        "filename":          doc_meta.get("filename", ""),
                        "description":       doc_meta.get("description", ""),
                        "summary":           doc_meta.get("summary", ""),
                    }

        hits = sorted(seen_docs.values(), key=lambda x: x["cosine_similarity"], reverse=True)
        return {"query": query, "results": hits, "total": len(hits)}

    except Exception as e:
        logger.warning("[KB] Search error: %s", e)
        return {"query": query, "results": [], "total": 0, "error": str(e)}


# ── Ticket similarity ─────────────────────────────────────────────────────────

def get_similar_docs_for_ticket(
    query: str,
    domain: Optional[str] = None,
    n_results: int = 5,
) -> dict:
    result = search_knowledge(query=query, n_results=n_results, domain=domain)
    if result["total"] < 2 and domain:
        broader = search_knowledge(query=query, n_results=n_results)
        if broader["total"] > result["total"]:
            result = broader
    return result


# ── RAG context for chat ──────────────────────────────────────────────────────

def get_rag_context(query: str, domain: Optional[str] = None, n_results: int = 3) -> str:
    # Check cache first
    cache_key = _rag_cache_key(query, domain, n_results)
    now_mono  = time.monotonic()
    with _rag_cache_lock:
        if cache_key in _rag_cache:
            cached_val, cached_ts = _rag_cache[cache_key]
            if now_mono - cached_ts < _RAG_CACHE_TTL:
                return cached_val

    try:
        collection = _get_collection()
        # Use timeout wrapper to avoid blocking the event loop on slow ChromaDB
        count = _chroma_with_timeout(collection.count)
        if count is None or count == 0:
            return ""

        result = search_knowledge(query=query, n_results=n_results, domain=domain)
        # 50% cosine similarity threshold — avoids off-topic runbook matches
        hits   = [r for r in result.get("results", []) if r["cosine_similarity"] >= 50]
        # If no hits with domain filter, retry without — document may have a different domain label
        if not hits and domain:
            result = search_knowledge(query=query, n_results=n_results, domain=None)
            hits   = [r for r in result.get("results", []) if r["cosine_similarity"] >= 50]
        if not hits:
            logger.info("[RAG] no hits above 50%% threshold for query: %.80s", query)
            return ""

        logger.info("[RAG] %d doc(s) matched — using:", len({h.get("doc_id") for h in hits}))
        for i, hit in enumerate(hits[:5], 1):
            src = hit.get("filename", "") or "(unknown)"
            logger.info("[RAG]   [%d] title=%r  domain=%s  sim=%s%%  src=%s",
                        i, hit["title"], hit["domain"], hit["cosine_similarity"], src)

        # Return chunks from every matched document.
        # Small docs (<= _MAX_CONTEXT_CHARS): include all chunks in order (full runbook).
        # Large docs (> _MAX_CONTEXT_CHARS): run a targeted similarity search within the doc
        # to get the most relevant chunks up to the cap, preserving their original order.
        _MAX_CONTEXT_CHARS = 20_000

        embedder        = _get_embedder()
        query_embedding = embedder.encode([query]).tolist()

        seen_docs = set()
        context   = "\n\n--- Knowledge Base ---\n"
        for hit in hits:
            doc_id = hit.get("doc_id", "")
            if doc_id and doc_id not in seen_docs:
                seen_docs.add(doc_id)
                try:
                    all_chunks = _chroma_with_timeout(
                        collection.get,
                        where={"doc_id": doc_id},
                        include=["documents", "metadatas"],
                    )
                    if all_chunks is None:
                        context += f"\n[{hit['title']}]\n{hit['content']}\n\n"
                        continue
                    if all_chunks and all_chunks["documents"]:
                        chunks_with_idx = list(zip(
                            all_chunks["documents"],
                            all_chunks["metadatas"],
                        ))
                        chunks_with_idx.sort(key=lambda x: x[1].get("chunk_index", 0))
                        doc_text_full = "\n\n".join(t for t, _ in chunks_with_idx)

                        if len(doc_text_full) <= _MAX_CONTEXT_CHARS:
                            context += f"\n[{hit['title']}]\n{doc_text_full}\n\n"
                        else:
                            # Large doc: similarity-search within this doc for the best chunks
                            doc_count = len(all_chunks["documents"])
                            try:
                                inner = _chroma_with_timeout(
                                    collection.query,
                                    query_embeddings=query_embedding,
                                    n_results=min(20, doc_count),
                                    where={"doc_id": doc_id},
                                    include=["documents", "metadatas", "distances"],
                                )
                            except Exception:
                                inner = None

                            if inner and inner["documents"] and inner["documents"][0]:
                                # Keep chunks in their original document order, up to cap
                                inner_by_idx = sorted(
                                    zip(inner["documents"][0], inner["metadatas"][0]),
                                    key=lambda x: x[1].get("chunk_index", 0),
                                )
                                selected, total_len = [], 0
                                for chunk_text, _ in inner_by_idx:
                                    if total_len + len(chunk_text) > _MAX_CONTEXT_CHARS:
                                        break
                                    selected.append(chunk_text)
                                    total_len += len(chunk_text)
                                context += f"\n[{hit['title']}]\n" + "\n\n".join(selected) + "\n\n"
                                logger.info(
                                    "[RAG] doc %s: %d chunks total → top %d by similarity (%d chars)",
                                    doc_id[:8], doc_count, len(selected), total_len,
                                )
                            else:
                                context += f"\n[{hit['title']}]\n{doc_text_full[:_MAX_CONTEXT_CHARS]}\n\n"
                except Exception:
                    context += f"\n[{hit['title']}]\n{hit['content']}\n\n"
        context += "\n--- End Knowledge Base ---\n"

        # Store in cache
        with _rag_cache_lock:
            if len(_rag_cache) >= _RAG_CACHE_MAX:
                # Evict oldest 10%
                sorted_items = sorted(_rag_cache.items(), key=lambda kv: kv[1][1])
                for k, _ in sorted_items[:max(1, _RAG_CACHE_MAX // 10)]:
                    _rag_cache.pop(k, None)
            _rag_cache[cache_key] = (context, time.monotonic())

        return context

    except Exception:
        return ""


# ── URL Auto-Refresh ─────────────────────────────────────────────────────────

import hashlib
import threading

def _content_hash(text: str) -> str:
    """Generate hash of content to detect changes."""
    return hashlib.md5(text.encode("utf-8"), usedforsecurity=False).hexdigest()


def _refresh_url_document(doc_id: str, meta_entry: dict) -> bool:
    """
    Re-fetch a URL document and re-index if content has changed.
    Returns True if content was updated, False if unchanged.
    """
    url = meta_entry.get("source_url") or meta_entry.get("filename", "")
    if not url or not url.startswith("http"):
        return False

    try:
        text = fetch_url_content(url)
        if not text.strip():
            logger.warning("[AutoRefresh] Empty content from URL (doc_id=%s)", doc_id)
            return False

        new_hash = _content_hash(text)
        old_hash = meta_entry.get("content_hash", "")

        if new_hash == old_hash:
            return False  # No change

        # Content changed — re-index
        logger.info("[AutoRefresh] '%s' content changed — re-indexing", meta_entry["title"])

        chunks     = _chunk_text(text)
        embedder   = _get_embedder()
        embeddings = embedder.encode(chunks).tolist()
        collection = _get_collection()

        # Delete old chunks and add new ones atomically under write lock
        with _collection_write_lock:
            try:
                old_results = collection.get(where={"doc_id": doc_id})
                if old_results["ids"]:
                    collection.delete(ids=old_results["ids"])
            except Exception:
                pass

            ids       = [f"{doc_id}_{i}" for i in range(len(chunks))]
            metadatas = [
                {"doc_id": doc_id, "domain": meta_entry["domain"],
                 "title": meta_entry["title"], "chunk_index": i, "source_url": url}
                for i in range(len(chunks))
            ]
            collection.add(ids=ids, embeddings=embeddings, documents=chunks, metadatas=metadatas)
        _clear_rag_cache()

        # Regenerate summary
        summary = _generate_summary(text[:8000], meta_entry["title"], meta_entry["domain"])

        # Update metadata
        meta = _load_meta()
        if doc_id in meta:
            meta[doc_id]["chunk_count"]   = len(chunks)
            meta[doc_id]["summary"]       = summary
            meta[doc_id]["content_hash"]  = new_hash
            meta[doc_id]["last_refreshed"] = datetime.utcnow().isoformat()
            _save_meta(meta)

        logger.info("[AutoRefresh] '%s' — %d chunks re-indexed", meta_entry["title"], len(chunks))
        return True

    except Exception as e:
        logger.warning("[AutoRefresh] Error for '%s': %s", meta_entry.get("title"), e)
        return False


_refresh_failure_counts: dict = {}
_REFRESH_BACKOFF_STEPS  = [120, 300, 900, 3600]  # 2m, 5m, 15m, 1h


def _auto_refresh_loop():
    """Background thread — checks all URL documents for changes, with per-URL failure backoff."""
    import time
    while True:
        time.sleep(120)
        try:
            meta = _load_meta()
            url_docs = {
                doc_id: entry for doc_id, entry in meta.items()
                if entry.get("source_type") == "url" or
                   entry.get("filename", "").startswith("http")
            }
            if not url_docs:
                continue
            logger.debug("[AutoRefresh] Checking %d URL documents for changes", len(url_docs))
            updated = 0
            for doc_id, entry in url_docs.items():
                fail_count = _refresh_failure_counts.get(doc_id, 0)
                backoff    = _REFRESH_BACKOFF_STEPS[min(fail_count, len(_REFRESH_BACKOFF_STEPS) - 1)]
                last_tried = entry.get("_last_refresh_attempt", 0)
                if time.time() - last_tried < backoff:
                    continue
                try:
                    refreshed = _refresh_url_document(doc_id, entry)
                    if refreshed:
                        updated += 1
                    _refresh_failure_counts[doc_id] = 0
                except Exception:
                    _refresh_failure_counts[doc_id] = fail_count + 1
                    logger.warning("[AutoRefresh] Failure #%d for doc %s — next retry in %ds",
                                   _refresh_failure_counts[doc_id], doc_id, backoff)
            logger.debug("[AutoRefresh] Complete: %d/%d documents updated", updated, len(url_docs))
        except Exception as e:
            logger.error("[AutoRefresh] Loop error: %s", e)


def start_url_refresh_scheduler():
    """Start the background URL refresh scheduler."""
    thread = threading.Thread(target=_auto_refresh_loop, daemon=True)
    thread.start()
    logger.info("[AutoRefresh] URL refresh scheduler started")


# ── List / Delete ─────────────────────────────────────────────────────────────

def list_documents(domain: Optional[str] = None) -> list:
    meta = _load_meta()
    docs = list(meta.values())
    if domain and domain not in ("all", ""):
        docs = [d for d in docs if d.get("domain") == domain]
    return sorted(docs, key=lambda x: x.get("created_at", ""), reverse=True)


def delete_document(doc_id: str) -> bool:
    meta = _load_meta()
    if doc_id not in meta:
        return False
    try:
        collection = _get_collection()
        with _collection_write_lock:
            results = collection.get(where={"doc_id": doc_id})
            if results["ids"]:
                collection.delete(ids=results["ids"])
    except Exception as e:
        logger.warning("[KB] Delete error for %s: %s", doc_id, e)
    del meta[doc_id]
    _save_meta(meta)
    _clear_rag_cache()
    logger.info("[KB] Deleted doc_id=%s", doc_id)
    return True